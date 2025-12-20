"""Text extraction utilities for multiple document formats.

Supports:
- PDF files (.pdf)
- Word documents (.docx)

Provides both function-based API (backward compatibility) and
class-based API (for dependency injection).
"""

import os


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text as string
    """
    from pdfminer.high_level import extract_text as pdf_extract_text
    return pdf_extract_text(file_path)


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a Word document (.docx).
    
    Args:
        file_path: Path to the Word document
        
    Returns:
        Extracted text as string
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for Word document support. "
            "Install it with: pip install python-docx"
        )

    doc = Document(file_path)

    # Extract text from paragraphs
    paragraphs = [para.text for para in doc.paragraphs]

    # Extract text from tables
    table_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            table_text.append(" | ".join(row_text))

    # Combine all text
    all_text = "\n".join(paragraphs + table_text)
    return all_text


def extract_text(file_path: str) -> str:
    """
    Extract text from a document, auto-detecting the format.
    
    Supported formats:
    - PDF (.pdf)
    - Word (.docx)
    
    Args:
        file_path: Path to the document
        
    Returns:
        Extracted text as string
        
    Raises:
        ValueError: If file format is not supported
        FileNotFoundError: If file does not exist
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(
            f"Unsupported file format: {ext}. "
            "Supported formats: .pdf, .docx"
        )


class TextExtractor:
    """Text extractor implementing TextExtractorProtocol.
    
    Provides a class-based interface for text extraction, enabling
    dependency injection and easier testing.
    
    Implements: core.protocols.TextExtractorProtocol
    
    Usage:
        extractor = TextExtractor()
        text = extractor.extract("document.pdf")
    """
    
    def extract(self, file_path: str) -> str:
        """Extract text from a document.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Extracted text as string
            
        Raises:
            FileNotFoundError: If file does not exist
            ValueError: If file format is not supported
        """
        return extract_text(file_path)


class MockTextExtractor:
    """Mock text extractor for testing.
    
    Implements: core.protocols.TextExtractorProtocol
    
    Usage:
        extractor = MockTextExtractor("This is test text")
        text = extractor.extract("any_path.pdf")  # Returns "This is test text"
    """
    
    def __init__(self, return_text: str = ""):
        """Initialize with text to return.
        
        Args:
            return_text: Text to return from extract()
        """
        self._return_text = return_text
        self.extract_called_with: list[str] = []  # Track calls for assertions
    
    def extract(self, file_path: str) -> str:
        """Return the configured text (for testing).
        
        Args:
            file_path: Path (stored for assertion)
            
        Returns:
            Configured return text
        """
        self.extract_called_with.append(file_path)
        return self._return_text

